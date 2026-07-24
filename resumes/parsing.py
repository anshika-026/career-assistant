"""
Text extraction utilities for uploaded resumes (PDF / DOCX).
Keeping this in its own module makes it easy to swap/extend later
(e.g. add OCR fallback for scanned PDFs) without touching views.
"""
import io
import logging

import pdfplumber
import docx

logger = logging.getLogger(__name__)


class UnsupportedFileTypeError(Exception):
    pass


class EmptyResumeError(Exception):
    """Raised when parsing succeeds but yields no usable text
    (e.g. a scanned/image-only PDF with no OCR)."""
    pass


def extract_text_from_pdf(file_obj) -> str:
    """Extract text from a PDF file object using pdfplumber."""
    text_chunks = []
    file_obj.seek(0)
    with pdfplumber.open(file_obj) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_chunks.append(page_text)
    return "\n".join(text_chunks).strip()


def extract_text_from_docx(file_obj) -> str:
    """Extract text from a DOCX file object using python-docx."""
    file_obj.seek(0)
    document = docx.Document(io.BytesIO(file_obj.read()))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

    # Also pull text out of tables (many resumes use table layouts)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())

    return "\n".join(paragraphs).strip()


def extract_resume_text(file_obj, file_type: str) -> str:
    """
    Dispatch to the right extractor based on file_type ('pdf' or 'docx').
    Raises UnsupportedFileTypeError / EmptyResumeError on failure so the
    caller (the view) can turn that into a clean API error response.
    """
    file_type = file_type.lower()

    try:
        if file_type == "pdf":
            text = extract_text_from_pdf(file_obj)
        elif file_type in ("docx", "doc"):
            text = extract_text_from_docx(file_obj)
        else:
            raise UnsupportedFileTypeError(
                f"Unsupported file type '{file_type}'. Only PDF and DOCX are supported."
            )
    except UnsupportedFileTypeError:
        raise
    except Exception as exc:
        logger.exception("Failed to parse resume file")
        raise EmptyResumeError(f"Could not parse file: {exc}") from exc

    if not text:
        raise EmptyResumeError(
            "No text could be extracted. The file may be a scanned image "
            "without a text layer (OCR not yet supported)."
        )

    return text
